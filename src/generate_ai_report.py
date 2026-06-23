from __future__ import annotations

import argparse
import base64
import json
import os
import re
import sys
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from openai import OpenAI, RateLimitError, APIError
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


# -----------------------------------------------------------------------------
# Defaults - can be overridden from command line
# -----------------------------------------------------------------------------

# Do not hardcode a date-specific ZIP path here. In production, run_all.py passes
# the AI ZIP path created for the current model run.
#
# Model can be overridden in this priority order:
# 1) --model CLI argument
# 2) SYNOPTICS_OPENAI_MODEL in .env
# 3) fallback below
DEFAULT_MODEL = "gpt-5.4-mini"


# -----------------------------------------------------------------------------
# Small helpers
# -----------------------------------------------------------------------------

def fmt(value: Any, ndigits: int = 1) -> str:
    if value is None:
        return "NA"
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        return f"{value:.{ndigits}f}"
    return str(value)


def safe_get(dct: dict, path: list[str], default: Any = None) -> Any:
    cur: Any = dct
    for key in path:
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur


def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9_\-]+", "_", value)
    return re.sub(r"_+", "_", value).strip("_") or "report"


def is_image_name(name: str) -> bool:
    return name.lower().endswith((".png", ".jpg", ".jpeg", ".webp"))


def mime_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".png":
        return "image/png"
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".webp":
        return "image/webp"
    return "application/octet-stream"


def image_to_data_url(path: Path) -> str:
    data = path.read_bytes()
    encoded = base64.b64encode(data).decode("ascii")
    return f"data:{mime_type(path)};base64,{encoded}"


def clean_model_output(text: str) -> str:
    """Remove common accidental wrappers without changing the report content."""
    text = text.strip()
    text = re.sub(r"^```(?:markdown|md|text)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    banned_starts = [
        "Nemohu přímo vytvořit",
        "Nemohu vytvořit",
        "Jako AI",
        "Jako jazykový model",
    ]
    lines = text.splitlines()
    while lines and any(lines[0].strip().startswith(s) for s in banned_starts):
        lines.pop(0)
    return "\n".join(lines).strip()


# -----------------------------------------------------------------------------
# briefing_context.json summarisation
# -----------------------------------------------------------------------------

def build_key_diagnostics(context: dict) -> str:
    """
    Automatically extracts the most important numerical diagnostics from
    briefing_context.json. No values are manually hard-coded into the prompt.
    """
    timesteps = context.get("timesteps", [])
    if not timesteps:
        return "No timesteps found in briefing_context.json."

    rows: list[dict[str, Any]] = []

    for step in timesteps:
        fxx = step.get("forecast_hour")
        valid = step.get("valid_time")
        cz = safe_get(step, ["features", "regions", "czechia"], {})
        cls = safe_get(step, ["commentary", "classification"], {})
        derived = safe_get(step, ["features", "derived"], {})

        rows.append({
            "forecast_hour": fxx,
            "valid_time": valid,
            "synoptic_type": cls.get("synoptic_type"),
            "weather_regime": cls.get("weather_regime"),
            "flow_type": cls.get("flow_type"),

            "cz_t850_mean": safe_get(cz, ["t850_c", "mean"]),
            "cz_t850_max": safe_get(cz, ["t850_c", "max"]),
            "cz_t850_ge20": safe_get(cz, ["t850_c", "area_fraction_percent", "ge_20c"]),

            "cz_mslp_mean": safe_get(cz, ["mslp_hpa", "mean"]),
            "cz_pwat_max": safe_get(cz, ["pwat_mm", "max"]),
            "cz_pwat_ge30": safe_get(cz, ["pwat_mm", "area_fraction_percent", "ge_30mm"]),

            "cz_cape_max": safe_get(cz, ["cape_jkg", "max"]),
            "cz_cape_ge1000": safe_get(cz, ["cape_jkg", "area_fraction_percent", "ge_1000jkg"]),

            "cz_cin_mean": safe_get(cz, ["cin_jkg", "mean"]),
            "cz_cin_max": safe_get(cz, ["cin_jkg", "max"]),
            "cz_cin_ge50": safe_get(cz, ["cin_jkg", "area_fraction_percent", "ge_50jkg"]),

            "cz_precip_max": safe_get(cz, ["precip_mm", "max"]),
            "cz_precip_ge1": safe_get(cz, ["precip_mm", "area_fraction_percent", "ge_1mm"]),

            "cz_jet250_max": safe_get(cz, ["jet250_speed_ms", "max"]),

            "eu_t850_ge20": derived.get("t850_area_gt_20c_percent") or derived.get("area_t850_ge_20c_percent"),
            "eu_precip_ge10": derived.get("precip_area_gt_10mm_percent"),
            "eu_jet_ge50": derived.get("jet250_area_gt_50ms_percent"),
            "eu_pwat_ge40": derived.get("pwat_area_gt_40mm_percent"),
            "eu_cape_ge1000": derived.get("cape_area_gt_1000jkg_percent"),
        })

    def max_row(key: str) -> dict[str, Any] | None:
        valid_rows = [r for r in rows if isinstance(r.get(key), (int, float))]
        return max(valid_rows, key=lambda r: r[key]) if valid_rows else None

    def min_row(key: str) -> dict[str, Any] | None:
        valid_rows = [r for r in rows if isinstance(r.get(key), (int, float))]
        return min(valid_rows, key=lambda r: r[key]) if valid_rows else None

    lines: list[str] = []
    lines.append("AUTOMATIC KEY DIAGNOSTICS FROM briefing_context.json")
    lines.append("")
    lines.append(
        f"Model: {context.get('model', 'NA')}; run: {context.get('run_time', 'NA')}; "
        f"domain: {context.get('domain', 'NA')}; focus region: {context.get('focus_region', 'NA')}."
    )
    lines.append("")
    lines.append("Main Czechia extremes across all forecast hours:")

    extremes = [
        ("Highest Czechia mean T850", "cz_t850_mean", "°C", 1, max_row),
        ("Highest Czechia max T850", "cz_t850_max", "°C", 1, max_row),
        ("Highest Czechia PWAT max", "cz_pwat_max", "mm", 1, max_row),
        ("Highest Czechia CAPE max", "cz_cape_max", "J/kg", 0, max_row),
        ("Highest Czechia mean CIN", "cz_cin_mean", "J/kg", 1, max_row),
        ("Highest Czechia precip max", "cz_precip_max", "mm", 1, max_row),
        ("Highest Czechia Jet250 max", "cz_jet250_max", "m/s", 1, max_row),
        ("Lowest Czechia mean MSLP", "cz_mslp_mean", "hPa", 1, min_row),
    ]
    for label, key, unit, ndigits, fn in extremes:
        r = fn(key)
        if r:
            lines.append(
                f"- {label}: {fmt(r[key], ndigits)} {unit} at +{r['forecast_hour']} h "
                f"({r['valid_time']})."
            )

    lines.append("")
    lines.append("Czechia time series by forecast hour:")
    lines.append(
        "fxx | valid_time | synoptic_type | weather_regime | "
        "T850 mean/max [°C] | T850 >=20 [%] | PWAT max [mm] | "
        "CAPE max [J/kg] | CAPE >=1000 [%] | CIN mean/max [J/kg] | "
        "CIN >=50 [%] | precip max [mm] | precip >=1 [%] | Jet250 max [m/s]"
    )

    for r in rows:
        lines.append(
            f"+{r['forecast_hour']} | {r['valid_time']} | {r['synoptic_type']} | {r['weather_regime']} | "
            f"{fmt(r['cz_t850_mean'])}/{fmt(r['cz_t850_max'])} | {fmt(r['cz_t850_ge20'])} | "
            f"{fmt(r['cz_pwat_max'])} | {fmt(r['cz_cape_max'], 0)} | {fmt(r['cz_cape_ge1000'])} | "
            f"{fmt(r['cz_cin_mean'])}/{fmt(r['cz_cin_max'])} | {fmt(r['cz_cin_ge50'])} | "
            f"{fmt(r['cz_precip_max'])} | {fmt(r['cz_precip_ge1'])} | {fmt(r['cz_jet250_max'])}"
        )

    return "\n".join(lines)


# -----------------------------------------------------------------------------
# ZIP reading and figure extraction
# -----------------------------------------------------------------------------

def read_zip_inputs(zip_path: Path) -> dict[str, Any]:
    print("Opening ZIP...")
    if not zip_path.exists():
        raise FileNotFoundError(f"ZIP not found: {zip_path}")

    with zipfile.ZipFile(zip_path, "r") as z:
        names = z.namelist()
        print(f"Files found: {len(names)}")

        context = None
        context_file = None
        template_text = ""
        instructions_text = ""
        prompt_text = ""

        for name in names:
            lower = name.lower()

            if lower.endswith("briefing_context.json"):
                context = json.loads(z.read(name).decode("utf-8"))
                context_file = name

            elif lower.endswith("diagnostics.json") and context is None:
                context = json.loads(z.read(name).decode("utf-8"))
                context_file = name

            elif lower.endswith("prompt.md"):
                prompt_text += f"\n\n--- PROMPT FILE: {name} ---\n"
                prompt_text += z.read(name).decode("utf-8", errors="replace")

            elif "template" in lower and lower.endswith((".md", ".txt")):
                template_text += f"\n\n--- TEMPLATE FILE: {name} ---\n"
                template_text += z.read(name).decode("utf-8", errors="replace")

            elif (
                "instruction" in lower
                or "report_specification" in lower
                or "style_guide" in lower
                or lower.endswith("start_here.md")
                or lower.endswith("readme.md")
            ) and lower.endswith((".md", ".txt")):
                instructions_text += f"\n\n--- INSTRUCTION FILE: {name} ---\n"
                instructions_text += z.read(name).decode("utf-8", errors="replace")

        if context is None:
            raise RuntimeError("Neither briefing_context.json nor diagnostics.json found inside ZIP.")

        combined_names = [
            n for n in names
            if "combined_figures/" in n.replace("\\", "/").lower() and is_image_name(n)
        ]

    return {
        "context": context,
        "context_file": context_file,
        "template_text": template_text,
        "instructions_text": instructions_text,
        "prompt_text": prompt_text,
        "combined_names": combined_names,
    }


def select_combined_figures(context: dict, combined_names: list[str], hours: list[int]) -> list[str]:
    if not combined_names:
        return []

    selected: list[str] = []
    for hour in hours:
        # Prefer filenames like f072, f120, etc.
        token = f"f{hour:03d}"
        matches = [n for n in combined_names if token in Path(n).name.lower()]
        if matches:
            selected.append(sorted(matches)[0])
            continue

        # Fallback: use context's combined_figure path.
        for step in context.get("timesteps", []):
            if step.get("forecast_hour") == hour and step.get("combined_figure"):
                cf = step["combined_figure"].replace("\\", "/")
                matches = [n for n in combined_names if n.replace("\\", "/").endswith(cf)]
                if matches:
                    selected.append(matches[0])
                break

    # Deduplicate while keeping order.
    seen = set()
    out = []
    for n in selected:
        if n not in seen:
            out.append(n)
            seen.add(n)
    return out




def parse_int_list(value: str | None) -> list[int]:
    if not value:
        return []
    return [int(item.strip()) for item in value.split(",") if item.strip()]


def representative_hours(hours: list[int], max_figures: int) -> list[int]:
    """Select representative forecast hours from the run_all --fxx-list.

    The selection is deterministic and derived from the actual forecast hours,
    not hard-coded to a specific model run/date. For a long forecast list this
    keeps the report compact while still showing the beginning, middle and end.
    """
    unique = sorted(set(int(h) for h in hours))
    if not unique:
        return []
    if max_figures <= 0 or len(unique) <= max_figures:
        return unique

    # Evenly spaced indices including first and last.
    n = len(unique)
    idxs = sorted({round(i * (n - 1) / (max_figures - 1)) for i in range(max_figures)})
    return [unique[i] for i in idxs]



def derive_max_figures(hours: list[int]) -> int:
    """Derive a compact figure count from the actual forecast configuration.

    This avoids a hard-coded DEFAULT_MAX_FIGURES. The goal is to show enough
    combined maps for the synoptic evolution without making the DOCX/PDF mostly
    figures.

    Rules:
    - 0-4 forecast times: use all.
    - 5-8 forecast times: use 4.
    - 9-14 forecast times: use 5.
    - 15+ forecast times: use 6.
    """
    unique = sorted(set(int(h) for h in hours))
    n = len(unique)
    if n <= 4:
        return n
    if n <= 8:
        return 4
    if n <= 14:
        return 5
    return 6

def resolve_figure_hours(context: dict, figure_hours_arg: str, fxx_list_arg: str | None, max_figures: int | None) -> list[int]:
    """Resolve figure hours from CLI arguments and/or briefing context.

    - --figure-hours auto: derive from --fxx-list if supplied by run_all.py.
      If --fxx-list is not supplied, derive from briefing_context.json timesteps.
    - --figure-hours all: use every forecast hour from --fxx-list/context.
    - --figure-hours 0,120,240: use explicit hours.
    """
    value = (figure_hours_arg or "auto").strip().lower()
    context_hours = [
        int(step["forecast_hour"])
        for step in context.get("timesteps", [])
        if isinstance(step.get("forecast_hour"), int)
    ]
    fxx_hours = parse_int_list(fxx_list_arg)
    base_hours = fxx_hours or context_hours
    effective_max_figures = max_figures if max_figures is not None else derive_max_figures(base_hours)

    if value == "all":
        return sorted(set(base_hours))

    if value == "auto":
        return representative_hours(base_hours, max_figures=effective_max_figures)

    return parse_int_list(figure_hours_arg)

def extract_selected_figures(zip_path: Path, selected_names: list[str], assets_dir: Path) -> list[Path]:
    assets_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    with zipfile.ZipFile(zip_path, "r") as z:
        for name in selected_names:
            out = assets_dir / Path(name).name
            out.write_bytes(z.read(name))
            paths.append(out)
    return paths


def build_figure_context(figure_paths: list[Path]) -> str:
    if not figure_paths:
        return "No combined figures were attached. Use the JSON diagnostics only."
    lines = ["Combined overview figures attached to this request:"]
    for p in figure_paths:
        lines.append(f"- {p.name}")
    lines.append("")
    lines.append(
        "Use these figures only for synoptic pattern recognition: ridge/trough structure, "
        "jet position, overlap of PWAT/CAPE/CIN/precipitation, and spatial context. "
        "Numerical values must still come from briefing_context.json or automatic key diagnostics."
    )
    return "\n".join(lines)


# -----------------------------------------------------------------------------
# Prompt and API call
# -----------------------------------------------------------------------------

def build_prompt(
    *,
    context: dict,
    key_diagnostics_text: str,
    figure_context: str,
    prompt_text: str,
    instructions_text: str,
    template_text: str,
    max_pages_hint: int,
) -> str:
    return f"""
You are a senior meteorological forecaster.

Create ONLY the final Czech synoptic forecast report text.

Important:
- Do not mention DOCX, files, attachments, limitations, ChatGPT, API, or formatting limitations.
- Do not say that you cannot create a file.
- The Python script will create the DOCX automatically.
- Return only the report content, starting with the report title.
- Do not include any preface, apology, explanation, or offer of further help.

Output style:
- Write in Czech.
- Be concise and professional.
- Target length: approximately {max_pages_hint} Word pages, not a long essay.
- Avoid repeating the same CAPE/CIN/PWAT explanation in every time block.
- Prefer grouped periods over one repetitive paragraph per forecast hour.
- Use concrete numbers from AUTOMATIC KEY DIAGNOSTICS and briefing_context.json.
- If you mention a numerical value, it must exist in the supplied data.
- Do not invent values.

Required report structure:
# Synoptický prognostický report GFS
## Stručné shrnutí
## Hlavní synoptický příběh
## Vývoj situace v čase
## Dopady na Evropu
## Dopady na Českou republiku
## Hlavní meteorologická rizika
## Prognostická důvěra

Do not create a separate section called "Interpretace map".
The Python script will insert the map figures visually into the final DOCX/PDF.
Use the attached figures for interpretation, but keep the final text concise.

=========================
AUTOMATIC KEY DIAGNOSTICS
=========================

{key_diagnostics_text}

=========================
COMBINED FIGURES CONTEXT
=========================

{figure_context}

=========================
PROMPT FILE
=========================

{prompt_text}

=========================
INSTRUCTIONS
=========================

{instructions_text}

=========================
TEMPLATE
=========================

{template_text}

=========================
BRIEFING CONTEXT JSON
=========================

{json.dumps(context, ensure_ascii=False, indent=2)}
"""


def call_openai(client: OpenAI, model: str, prompt: str, figure_paths: list[Path], use_vision: bool) -> str:
    content: list[dict[str, Any]] = [{"type": "input_text", "text": prompt}]

    if use_vision:
        for path in figure_paths:
            content.append({
                "type": "input_image",
                "image_url": image_to_data_url(path),
                "detail": "low",
            })

    print(f"Calling {model}...")
    print(f"Vision images sent: {len(figure_paths) if use_vision else 0}")

    try:
        response = client.responses.create(
            model=model,
            input=[{"role": "user", "content": content}],
            max_output_tokens=8000,
        )
    except RateLimitError:
        print("OPENAI QUOTA/RATE LIMIT ERROR: check API billing, credits, and limits.")
        raise
    except APIError:
        print("OPENAI API ERROR: request reached the API but failed.")
        raise

    return clean_model_output(response.output_text)


# -----------------------------------------------------------------------------
# DOCX formatting
# -----------------------------------------------------------------------------

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_document_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11, "4F81BD"),
    ]:
        st = styles[name]
        st.font.name = "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True


def configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Synoptický prognostický report GFS"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Automaticky generovaný meteorologický briefing"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_title_page(doc: Document, context: dict, model: str, zip_path: Path) -> None:
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Synoptický prognostický report GFS")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Evropa a Česká republika")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    meta = [
        ("Model", context.get("model", "GFS")),
        ("Běh modelu", context.get("run_time", "NA")),
        ("Doména", context.get("domain", "Europe")),
        ("Zájmový region", context.get("focus_region", "Česká republika")),
        ("AI model", model),
        ("Zdrojový balíček", zip_path.name),
    ]

    for i, (k, v) in enumerate(meta):
        set_cell_text(table.cell(i, 0), k, bold=True)
        set_cell_text(table.cell(i, 1), str(v))
        set_cell_shading(table.cell(i, 0), "D9EAF7")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(f"Vygenerováno: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def add_key_diagnostics_table(doc: Document, context: dict) -> None:
    timesteps = context.get("timesteps", [])
    if not timesteps:
        return

    doc.add_heading("Přehled klíčových diagnostik pro ČR", level=1)

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Fxx", "Valid UTC", "T850 mean/max", "PWAT max", "CAPE max", "CIN mean/max", "Srážky max", "Jet250 max"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        set_cell_shading(table.cell(0, i), "1F4E79")
        for run in table.cell(0, i).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for step in timesteps:
        cz = safe_get(step, ["features", "regions", "czechia"], {})
        cells = table.add_row().cells
        values = [
            f"+{step.get('forecast_hour')}",
            str(step.get("valid_time", "")),
            f"{fmt(safe_get(cz, ['t850_c', 'mean']))}/{fmt(safe_get(cz, ['t850_c', 'max']))} °C",
            f"{fmt(safe_get(cz, ['pwat_mm', 'max']))} mm",
            f"{fmt(safe_get(cz, ['cape_jkg', 'max'], 0), 0)} J/kg",
            f"{fmt(safe_get(cz, ['cin_jkg', 'mean']))}/{fmt(safe_get(cz, ['cin_jkg', 'max']))} J/kg",
            f"{fmt(safe_get(cz, ['precip_mm', 'max']))} mm",
            f"{fmt(safe_get(cz, ['jet250_speed_ms', 'max']))} m/s",
        ]
        for i, v in enumerate(values):
            set_cell_text(cells[i], v)

    doc.add_paragraph("")


def add_markdown_text(doc: Document, text: str) -> None:
    """Simple Markdown-ish to DOCX converter for headings, bullets and paragraphs."""
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            # Avoid duplicate title page look; first H1 becomes heading 1.
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s+", "", line).strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.08
            p.add_run(line.strip())


def add_figures_appendix(doc: Document, figure_paths: list[Path]) -> None:
    """Insert combined figures cleanly without technical filenames or blank pages."""
    if not figure_paths:
        return

    doc.add_page_break()
    doc.add_heading("Mapové podklady", level=1)

    for path in figure_paths:
        try:
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.paragraph_format.space_before = Pt(4)
            pic_p.paragraph_format.space_after = Pt(8)
            pic_p.add_run().add_picture(str(path), width=Inches(6.6))
        except Exception as exc:
            doc.add_paragraph(f"[Obrázek se nepodařilo vložit: {exc}]")

        # Important: no page break between images. Word/PDF decides naturally
        # where the next image fits, preventing empty pages between figures.


def export_pdf_from_docx(docx_path: Path, pdf_path: Path) -> bool:
    """
    Export DOCX to PDF.

    On Windows this first tries docx2pdf, which uses Microsoft Word.
    If that is not available, it tries LibreOffice / soffice when present.
    Returns True when PDF was created.
    """
    # 1) Preferred on Windows with Microsoft Word installed.
    try:
        from docx2pdf import convert  # type: ignore

        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"DOCX2PDF export unavailable/failed: {exc}")

    # 2) Fallback: LibreOffice headless export.
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(pdf_path.parent),
                    str(docx_path),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            generated = docx_path.with_suffix(".pdf")
            if generated.exists() and generated != pdf_path:
                generated.replace(pdf_path)
            return pdf_path.exists()
        except Exception as exc:
            print(f"LibreOffice PDF export failed: {exc}")

    print("PDF export skipped. Install docx2pdf + Microsoft Word, or LibreOffice.")
    return False


def create_docx_report(
    *,
    output_path: Path,
    report_text: str,
    context: dict,
    model: str,
    zip_path: Path,
    figure_paths: list[Path],
) -> None:
    doc = Document()
    set_document_styles(doc)
    configure_sections(doc)
    add_title_page(doc, context, model, zip_path)
    add_key_diagnostics_table(doc, context)
    add_markdown_text(doc, report_text)
    add_figures_appendix(doc, figure_paths)
    doc.save(output_path)


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate AI synoptic DOCX/PDF report from Synoptics AI ZIP.")
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Project root. Defaults to the parent directory of src/.",
    )
    parser.add_argument(
        "--zip",
        dest="zip_path",
        type=Path,
        required=True,
        help="Path to the AI input ZIP created for the current run. This is normally passed by run_all.py.",
    )
    parser.add_argument("--model", default=None, help="OpenAI model. If omitted, uses SYNOPTICS_OPENAI_MODEL or DEFAULT_MODEL.")
    parser.add_argument(
        "--fxx-list",
        default=None,
        help="Forecast hours from run_all.py, e.g. 0,6,12,24,48,72,96,120,144,168,192,216,240.",
    )
    parser.add_argument(
        "--figure-hours",
        default="auto",
        help="auto, all, or comma-separated forecast hours. 'auto' derives representative hours from --fxx-list/context.",
    )
    parser.add_argument(
        "--max-figures",
        type=int,
        default=None,
        help=(
            "Maximum number of combined figures when --figure-hours auto is used. "
            "If omitted, it is derived automatically from --fxx-list/context. "
            "Set explicitly only if you want to force a limit."
        ),
    )
    parser.add_argument("--no-vision", action="store_true", help="Do not send figures to GPT; still insert them into DOCX.")
    parser.add_argument("--max-pages-hint", type=int, default=4)
    parser.add_argument("--no-pdf", action="store_true", help="Do not export DOCX to PDF.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    project_root = args.project_root
    zip_path = args.zip_path if args.zip_path.is_absolute() else project_root / args.zip_path
    output_dir = zip_path.parent
    assets_dir = output_dir / "gpt_report_assets"

    load_dotenv(project_root / ".env")
    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError(f"OPENAI_API_KEY not found. Check {project_root / '.env'}")

    model = args.model or os.getenv("SYNOPTICS_OPENAI_MODEL") or DEFAULT_MODEL

    client = OpenAI()

    loaded = read_zip_inputs(zip_path)
    context = loaded["context"]
    context_file = loaded["context_file"]

    figure_hours = resolve_figure_hours(
        context=context,
        figure_hours_arg=args.figure_hours,
        fxx_list_arg=args.fxx_list,
        max_figures=args.max_figures,
    )
    selected_names = select_combined_figures(context, loaded["combined_names"], figure_hours)
    figure_paths = extract_selected_figures(zip_path, selected_names, assets_dir)

    key_diagnostics_text = build_key_diagnostics(context)
    figure_context = build_figure_context(figure_paths)

    print(f"Context loaded from: {context_file}")
    print(f"OpenAI model: {model}")
    print(f"Figure selection mode: {args.figure_hours}")
    print(f"Max figures override: {args.max_figures if args.max_figures is not None else 'auto-derived'}")
    print("Template length:", len(loaded["template_text"]))
    print("Instructions length:", len(loaded["instructions_text"]))
    print("Prompt length:", len(loaded["prompt_text"]))
    print("Key diagnostics length:", len(key_diagnostics_text))
    print("Combined figures selected:", len(figure_paths))
    for p in figure_paths:
        print(" -", p.name)

    prompt = build_prompt(
        context=context,
        key_diagnostics_text=key_diagnostics_text,
        figure_context=figure_context,
        prompt_text=loaded["prompt_text"],
        instructions_text=loaded["instructions_text"],
        template_text=loaded["template_text"],
        max_pages_hint=args.max_pages_hint,
    )

    # Save reproducibility artefacts.
    (output_dir / "gpt_prompt_used.md").write_text(prompt, encoding="utf-8")
    (output_dir / "gpt_key_diagnostics.md").write_text(key_diagnostics_text, encoding="utf-8")

    report_text = call_openai(
        client=client,
        model=model,
        prompt=prompt,
        figure_paths=figure_paths,
        use_vision=not args.no_vision,
    )

    md_path = output_dir / "gpt_report.md"
    docx_path = output_dir / "gpt_report.docx"

    md_path.write_text(report_text, encoding="utf-8")
    print(f"Saved: {md_path}")

    create_docx_report(
        output_path=docx_path,
        report_text=report_text,
        context=context,
        model=model,
        zip_path=zip_path,
        figure_paths=figure_paths,
    )
    print(f"Saved: {docx_path}")

    if not args.no_pdf:
        pdf_path = output_dir / "gpt_report.pdf"
        if export_pdf_from_docx(docx_path, pdf_path):
            print(f"Saved: {pdf_path}")
        else:
            print("PDF was not created.")

    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
